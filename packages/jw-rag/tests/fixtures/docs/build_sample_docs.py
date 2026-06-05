"""Genera un .docx de prueba simulando un 'Programa de Circuito' breve.

Contenido sintético, sin texto JW real, para evitar copyright en tests.

Para regenerar:
    cd packages/jw-rag/tests/fixtures/docs
    uv run --with python-docx python build_sample_docs.py

Requires: python-docx (dep dev only).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

HERE = Path(__file__).parent
OUTPUT = HERE / "programa_circuito.docx"


def main() -> None:
    doc = Document()
    doc.add_heading("Programa de Circuito — Sample Fixture", level=1)
    doc.add_paragraph(
        "Documento sintético para testing. NO contiene contenido JW real."
    )
    doc.add_heading("Reunión 1", level=2)
    doc.add_paragraph(
        "Discurso público: Lorem ipsum dolor sit amet, consectetur adipiscing elit."
    )
    doc.add_paragraph(
        "Estudio de la Atalaya: Sed do eiusmod tempor incididunt ut labore."
    )
    doc.add_heading("Reunión 2", level=2)
    doc.add_paragraph(
        "Vida y Ministerio Cristianos: Ut enim ad minim veniam, quis nostrud."
    )
    table = doc.add_table(rows=3, cols=2)
    table.style = "Light Grid"
    table.rows[0].cells[0].text = "Hora"
    table.rows[0].cells[1].text = "Parte"
    table.rows[1].cells[0].text = "10:00"
    table.rows[1].cells[1].text = "Cántico y oración"
    table.rows[2].cells[0].text = "10:15"
    table.rows[2].cells[1].text = "Discurso público"
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT} ({OUTPUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
