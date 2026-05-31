"""DOCX exporter via python-docx.

Uses python-docx's programmatic API directly (no template — DOCX templating
adds complexity without value at our structure level).
"""

from __future__ import annotations

from pathlib import Path

from jw_core.exporters.errors import MissingDependencyError
from jw_core.exporters.ir import CitationIR, StudySheet


def export_docx(sheet: StudySheet, *, out: Path) -> Path:
    """Render `sheet` as DOCX and write it to `out`. Returns `out`.

    Requires the [docx] extra. Raises `MissingDependencyError` otherwise.
    """

    try:
        from docx import Document  # noqa: PLC0415  (lazy)
        from docx.oxml.ns import qn  # noqa: PLC0415
        from docx.oxml import OxmlElement  # noqa: PLC0415
    except ImportError as exc:
        raise MissingDependencyError(
            "python-docx is required for DOCX export. "
            "Install with: pip install 'jw-core[docx]'"
        ) from exc

    doc = Document()

    # Title
    doc.add_heading(sheet.title, level=0)
    if sheet.subtitle:
        p = doc.add_paragraph()
        run = p.add_run(sheet.subtitle)
        run.italic = True

    # Sections
    for section in sheet.sections:
        doc.add_heading(section.heading, level=2)
        doc.add_paragraph(section.body)

        if section.excerpt:
            p = doc.add_paragraph(section.excerpt)
            p.style = doc.styles["Intense Quote"]

        for cite in section.citations:
            _add_citation_paragraph(doc, cite, qn, OxmlElement)

    if sheet.footer_note:
        doc.add_paragraph()
        sep = doc.add_paragraph("—" * 30)
        sep.alignment = 1  # center
        p = doc.add_paragraph()
        run = p.add_run(sheet.footer_note)
        run.italic = True
        run.font.size = run.font.size  # no-op to anchor formatting

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def _add_citation_paragraph(doc, cite: CitationIR, qn, OxmlElement) -> None:
    """Add a paragraph holding a hyperlink to the citation URL."""

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = p.paragraph_format.left_indent  # no-op
    label = cite.short_label or cite.title or cite.url

    # Add a real hyperlink relationship.
    part = p.part
    rid = part.relate_to(
        cite.url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0A3A6A")
    r_pr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    new_run.append(r_pr)

    t = OxmlElement("w:t")
    t.text = f"  • {label}"
    new_run.append(t)
    hyperlink.append(new_run)
    p._p.append(hyperlink)
